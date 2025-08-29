from __future__ import annotations

import re
import time
from pathlib import Path
from typing import List, Tuple, Dict, Any
from urllib.parse import urlparse

import requests
from docx import Document  # type: ignore


def extract_links_from_docx(docx_path: str) -> List[Dict[str, str]]:
    doc = Document(docx_path)
    links: List[Dict[str, str]] = []

    # Paragraph hyperlinks
    for paragraph in doc.paragraphs:
        if paragraph._element.xpath('.//w:hyperlink'):
            for hyperlink in paragraph._element.xpath('.//w:hyperlink'):
                rel_id = hyperlink.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                if not rel_id:
                    continue
                rels = doc.part.rels
                if rel_id not in rels:
                    continue
                url = rels[rel_id].target_ref
                text_elems = hyperlink.xpath('.//w:t')
                text = ''.join([e.text for e in text_elems if getattr(e, 'text', None)])
                text = (text or '').strip()
                context = (paragraph.text or '').strip()
                prefix = ""
                if text and context:
                    idx = context.find(text)
                    if idx >= 0:
                        prefix = context[:idx].strip()
                if url and (text or context):
                    links.append({'url': url, 'text': text, 'context': context, 'prefix': prefix, 'type': 'hyperlink'})

    # Runs (fallback)
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run._element.xpath('.//w:hyperlink'):
                for hyperlink in run._element.xpath('.//w:hyperlink'):
                    rel_id = hyperlink.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                    if not rel_id:
                        continue
                    rels = doc.part.rels
                    if rel_id not in rels:
                        continue
                    url = rels[rel_id].target_ref
                    text = (run.text or '').strip()
                    context = (paragraph.text or '').strip()
                    prefix = ""
                    if text and context:
                        idx = context.find(text)
                        if idx >= 0:
                            prefix = context[:idx].strip()
                    if url and (text or context):
                        links.append({'url': url, 'text': text, 'context': context, 'prefix': prefix, 'type': 'hyperlink_run'})

    # Headers
    for section in doc.sections:
        for header in section.header.paragraphs:
            if header._element.xpath('.//w:hyperlink'):
                for hyperlink in header._element.xpath('.//w:hyperlink'):
                    rel_id = hyperlink.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                    if not rel_id:
                        continue
                    rels = doc.part.rels
                    if rel_id not in rels:
                        continue
                    url = rels[rel_id].target_ref
                    text = (header.text or '').strip()
                    prefix = ""
                    if text:
                        # For headers we treat the whole header as context; no precise runs
                        # We cannot reliably find prefix; best-effort none or before text occurrence
                        idx = text.find(text)
                        prefix = text[:idx].strip() if idx > 0 else ""
                    if url and text:
                        links.append({'url': url, 'text': text, 'context': text, 'prefix': prefix, 'type': 'header'})

    # Deduplicate by URL
    seen = set()
    uniq: List[Dict[str, str]] = []
    for l in links:
        if l['url'] in seen:
            continue
        seen.add(l['url'])
        uniq.append(l)
    return uniq


def is_valid_url(url: str) -> bool:
    try:
        parts = urlparse(url)
        return bool(parts.scheme and parts.netloc)
    except Exception:
        return False


def get_filename_from_url(url: str, text: str | None = None) -> str:
    if text and text.strip():
        name = re.sub(r'[<>:"/\\|?*]', '_', text.strip())
        name = re.sub(r'\s+', '_', name)[:80]
    else:
        parts = urlparse(url)
        name = parts.netloc
        if parts.path:
            tail = Path(parts.path).name
            name = tail or (parts.path.replace('/', '_') or parts.netloc)
    if not name:
        name = 'downloaded_file'
    if '.' not in name:
        name += '.pdf'
    return name


VOLUME_RE = re.compile(r"Volume\s+(\d+)", re.IGNORECASE)

def suggest_filename_from_link(link: Dict[str, Any]) -> str | None:
    """
    If link text/context matches legal style like:
      "Volume 1573, Pages 131; 136; ..." or just "131" (with volume in context),
    build names:
      - Volume_<vol>_Page_<p>.pdf for single page
      - Volume_<vol>_Pages_<p1>_<p2>_... .pdf for multiple pages
    Returns None if no pattern found.
    """
    text = (link.get('text') or '').strip()
    context = (link.get('context') or '').strip()

    # Find volume in text or context
    m = VOLUME_RE.search(text) or VOLUME_RE.search(context)
    if not m:
        return None
    vol = m.group(1)

    # Collect numbers from the link text (prefer) else from context after 'Pages'
    nums = [n for n in re.findall(r"\d+", text)]
    # Exclude the volume number if present among nums
    nums = [n for n in nums if n != vol]
    if not nums:
        # Try to find after 'Pages' in context
        m_pages = re.search(r"Pages?\s+(.+)$", context, flags=re.IGNORECASE)
        if m_pages:
            nums = re.findall(r"\d+", m_pages.group(1))

    if not nums:
        return None

    if len(nums) == 1:
        return f"Volume_{vol}_Page_{nums[0]}.pdf"
    else:
        return f"Volume_{vol}_Pages_{'_'.join(nums)}.pdf"


def download_file(url: str, output_dir: Path, filename: str | None = None, timeout: int = 30):
    try:
        headers = {'User-Agent': 'Mozilla/5.0'}
        resp = requests.get(url, headers=headers, timeout=timeout, allow_redirects=True, stream=True)
        resp.raise_for_status()
        if not filename:
            cd = resp.headers.get('content-disposition')
            if cd and 'filename=' in cd:
                import cgi
                _, params = cgi.parse_header(cd)
                filename = params.get('filename')
        if not filename:
            filename = get_filename_from_url(url)
        filename = get_filename_from_url(url, filename)
        out = output_dir / filename
        counter = 1
        base = out
        while out.exists():
            out = output_dir / f"{base.stem}_{counter}{base.suffix}"
            counter += 1
        with open(out, 'wb') as f:
            for chunk in resp.iter_content(chunk_size=8192):
                if chunk:
                    f.write(chunk)
        return out, out.stat().st_size, None
    except requests.exceptions.Timeout:
        return None, 0, f"Timeout after {timeout} seconds"
    except requests.exceptions.ConnectionError:
        return None, 0, "Connection error"
    except requests.exceptions.HTTPError as e:
        return None, 0, f"HTTP error: {e.response.status_code}"
    except Exception as e:
        return None, 0, f"Unexpected error: {e}"


def save_error_report(error_report_file: Path, failed_downloads: List[Dict[str, Any]], docx_path: Path) -> None:
    with open(error_report_file, 'w', encoding='utf-8') as f:
        f.write("FAILED DOWNLOADS REPORT\n")
        f.write("=" * 30 + "\n\n")
        f.write(f"Source document: {docx_path}\n")
        f.write(f"Total failed downloads: {len(failed_downloads)}\n")
        f.write(f"Generated: {time.strftime('%Y-%m-%d %H:%M:%S')}\n\n")
        for i, failed in enumerate(failed_downloads, 1):
            f.write(f"{i}. Link #{failed['link_number']}: {failed['text']}\n")
            f.write(f"   URL: {failed['url']}\n")
            f.write(f"   Error: {failed['error']}\n\n")


def save_retry_script(retry_script_file: Path, failed_downloads: List[Dict[str, Any]], output_dir: Path) -> None:
    with open(retry_script_file, 'w', encoding='utf-8') as f:
        f.write("#!/usr/bin/env python3\n")
        f.write('"""Script to retry failed downloads from word-links"""\n\n')
        f.write("import requests\n")
        f.write("from pathlib import Path\n")
        f.write("import time\n\n")
        f.write(f"output_dir = Path(r\"{output_dir}\")\n\n")
        f.write("failed_downloads = [\n")
        for fail in failed_downloads:
            f.write("    {\n")
            f.write(f"        'url': r'{fail['url']}',\n")
            f.write(f"        'text': r'{fail['text']}',\n")
            f.write(f"        'error': r'{fail['error']}',\n")
            f.write(f"        'link_number': {fail['link_number']}\n")
            f.write("    },\n")
        f.write("]\n\n")
        f.write("def retry(url, text):\n")
        f.write("    headers={'User-Agent':'Mozilla/5.0'}\n")
        f.write("    r=requests.get(url, headers=headers, timeout=60, stream=True)\n")
        f.write("    r.raise_for_status()\n")
        f.write("    name=text.replace(' ','_')[:80]\n")
        f.write("    if '.' not in name: name += '.pdf'\n")
        f.write("    out=output_dir/name\n")
        f.write("    c=1; base=out\n")
        f.write("    while out.exists(): out=output_dir/f'{base.stem}_{c}{base.suffix}'; c+=1\n")
        f.write("    with open(out,'wb') as f:\n")
        f.write("        for ch in r.iter_content(8192):\n")
        f.write("            if ch: f.write(ch)\n")
        f.write("    return True\n\n")
        f.write("if __name__=='__main__':\n")
        f.write("    ok=0\n")
        f.write("    for it in failed_downloads:\n")
        f.write("        try:\n")
        f.write("            if retry(it['url'], it['text']): ok+=1\n")
        f.write("        except Exception as e: print('fail',e)\n")
        f.write("    print(f'Success {ok}/{len(failed_downloads)}')\n")


