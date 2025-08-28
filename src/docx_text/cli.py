import argparse
import sys
from pathlib import Path
from typing import List


def extract_docx_text(path: Path) -> str:
    from docx import Document  # type: ignore

    doc = Document(str(path))
    parts: List[str] = []
    # Paragraphs
    for p in doc.paragraphs:
        if p.text:
            parts.append(p.text)
    # Tables
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                t = cell.text.strip()
                if t:
                    row_text.append(t)
            if row_text:
                parts.append("\t".join(row_text))
    return "\n".join(parts) + "\n"


def build_arg_parser() -> argparse.ArgumentParser:
    p = argparse.ArgumentParser(
        prog="docx-to-text",
        description="Extract plain text from .docx files",
    )
    p.add_argument("inputdir", help="Directory containing .docx files or a single .docx path")
    p.add_argument("outputdir", help="Directory to write .txt outputs")
    p.add_argument("--recursive", action="store_true", help="Recurse into subdirectories")
    return p


def main(argv: List[str] | None = None) -> int:
    argv = sys.argv[1:] if argv is None else argv
    args = build_arg_parser().parse_args(argv)

    in_path = Path(args.inputdir)
    out_dir = Path(args.outputdir)
    out_dir.mkdir(parents=True, exist_ok=True)

    candidates: List[Path] = []
    if in_path.is_file() and in_path.suffix.lower() == ".docx":
        candidates = [in_path]
    elif in_path.is_dir():
        it = in_path.rglob("*.docx") if args.recursive else in_path.glob("*.docx")
        candidates = sorted(p for p in it if p.is_file())
    else:
        print("No .docx inputs found", file=sys.stderr)
        return 2

    for src in candidates:
        rel = src.relative_to(in_path) if in_path.is_dir() else Path(src.name)
        dst = out_dir / rel.with_suffix(".txt")
        dst.parent.mkdir(parents=True, exist_ok=True)
        try:
            text = extract_docx_text(src)
            dst.write_text(text, encoding="utf-8")
            print(f"Wrote {dst}")
        except Exception as e:
            print(f"Failed {src}: {e}", file=sys.stderr)
    return 0


